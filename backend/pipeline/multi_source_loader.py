"""Multi-format, multi-source document ingestion.

``load_document`` dispatches on file extension (.pdf / .docx / .txt), then
always chunks + vector-indexes the content, and finally runs the
source-type-specific processing:

- ``tender``            -> :func:`backend.pipeline.extraction.extract_requirements`
- ``change_request``    -> :func:`backend.pipeline.extraction.process_change_request`
- ``inspection_report`` -> LLM extraction of :class:`InspectionFinding` rows
  (same forced structured-output pattern via :mod:`backend.llm`; skipped with
  a warning when no LLM provider is configured)
- ``meeting_notes``     -> retrieval indexing only

Every ingestion records an :class:`IngestedDocument` row.

``pdf_loader``/``chunking``/``embedding`` are imported lazily inside
``load_document`` so importing this module stays cheap (no torch/chroma) and
tests can substitute them.
"""
from __future__ import annotations

import logging
import os
from typing import Any, Literal

import pydantic
from sqlalchemy.orm import Session

from backend import config, llm  # noqa: F401 — config kept for test monkeypatching
from backend.db.models import IngestedDocument, InspectionFinding
from backend.pipeline import extraction

logger = logging.getLogger(__name__)

SOURCE_TYPES = ("tender", "meeting_notes", "inspection_report", "change_request")

FINDINGS_TOOL_NAME = "record_findings"

FINDINGS_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "findings": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "location": {
                        "type": "string",
                        "description": "Where the defect was observed.",
                    },
                    "defect_description": {
                        "type": "string",
                        "description": "What was observed, as stated.",
                    },
                    "severity": {
                        "type": "string",
                        "enum": ["low", "medium", "high"],
                    },
                    "source_page": {
                        "type": "integer",
                        "description": "Page number the finding came from.",
                    },
                },
                "required": ["location", "defect_description", "severity"],
            },
        }
    },
    "required": ["findings"],
}

FINDINGS_SYSTEM_PROMPT = (
    "You are a construction quality inspector's assistant. You are given "
    "excerpts of a site inspection report, each tagged with its page number. "
    "Extract ONLY the findings explicitly stated in the text (location, "
    "defect description, severity). Never invent findings or infer details "
    "that are not stated. If there are no findings, record an empty list. "
    "Always respond via the record_findings tool."
)


class InspectionFindingModel(pydantic.BaseModel):
    """Validated shape of a single inspection finding."""

    location: str
    defect_description: str
    severity: Literal["low", "medium", "high"]
    source_page: int = 0


def _read_pages(file_path: str) -> list[dict]:
    """Dispatch on extension and return pdf_loader-shaped page dicts."""
    basename = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        from backend.pipeline.pdf_loader import extract_pdf

        return extract_pdf(file_path)

    if ext == ".docx":
        import docx  # python-docx; lazy so .txt/.pdf paths don't require it

        document = docx.Document(file_path)
        text = "\n".join(p.text for p in document.paragraphs)
        # Tenders typically carry the BOQ (grades/quantities) in tables, so
        # dropping them would lose the requirements entirely.
        tables = [
            [[cell.text for cell in row.cells] for row in table.rows]
            for table in document.tables
        ]
        return [
            {
                "page_number": 1,
                "text": text,
                "tables": tables,
                "source_file": basename,
            }
        ]

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as handle:
            text = handle.read()
        return [
            {
                "page_number": 1,
                "text": text,
                "tables": [],
                "source_file": basename,
            }
        ]

    raise ValueError(
        f"Unsupported file extension {ext!r} — expected .pdf, .docx or .txt"
    )


def _extract_inspection_findings(
    chunks: list[dict], project_id: str, db_session: Session
) -> list[InspectionFinding]:
    """LLM-extract InspectionFinding rows via forced structured output.

    Skips silently (with a warning log) when no LLM provider is configured.
    """
    if not llm.is_configured():
        logger.warning(
            "No LLM provider configured (ANTHROPIC_API_KEY / GEMINI_API_KEY "
            "not set) — skipping inspection-finding extraction for project %s.",
            project_id,
        )
        return []

    text_chunks = [c for c in chunks if c.get("chunk_type") != "table"]
    if not text_chunks:
        return []

    rows: list[InspectionFinding] = []

    for start in range(0, len(text_chunks), extraction.BATCH_SIZE):
        batch = text_chunks[start : start + extraction.BATCH_SIZE]
        default_file = batch[0].get("source_file", "")
        parts = [
            f"[page {c.get('page_number', 0)}]\n{c.get('text', '')}" for c in batch
        ]
        payload = llm.extract_json(
            FINDINGS_SYSTEM_PROMPT,
            (
                "Extract the explicitly stated inspection findings "
                "from the following report excerpts:\n\n"
                + "\n\n---\n\n".join(parts)
            ),
            FINDINGS_SCHEMA,
            tool_name=FINDINGS_TOOL_NAME,
        )
        if payload is None:
            logger.warning(
                "LLM findings extraction returned nothing for project %s (batch %d)",
                project_id,
                start // extraction.BATCH_SIZE,
            )
            continue

        raw = payload.get("findings", [])
        items: list[dict] = raw if isinstance(raw, list) else []

        for item in items:
            try:
                model_item = InspectionFindingModel.model_validate(item)
            except pydantic.ValidationError as exc:
                logger.warning("Skipping invalid inspection finding %r: %s", item, exc)
                continue
            rows.append(
                InspectionFinding(
                    project_id=project_id,
                    source_file=default_file,
                    location=model_item.location,
                    defect_description=model_item.defect_description,
                    severity=model_item.severity,
                    source_page=model_item.source_page,
                )
            )

    if rows:
        db_session.add_all(rows)
        db_session.commit()
    logger.info(
        "Recorded %d inspection finding(s) for project %s", len(rows), project_id
    )
    return rows


def load_document(
    file_path: str,
    project_id: str,
    source_type: str,
    db_session: Session,
    original_filename: str | None = None,
) -> dict:
    """Ingest one document end-to-end and record an IngestedDocument row.

    ``original_filename`` is the user-facing name of the document (e.g. the
    uploaded file's name) when ``file_path`` points at a staging copy such as
    a temp file; it is used for the IngestedDocument row, chunk/citation
    ``source_file`` metadata and the returned ``filename``.

    Returns the contract dict::

        {"project_id", "filename", "source_type", "pages_processed": int,
         "chunks_indexed": int, "tables_found": int,
         "requirements_extracted": int}
    """
    if source_type not in SOURCE_TYPES:
        raise ValueError(
            f"Unknown source_type {source_type!r} — expected one of {SOURCE_TYPES}"
        )

    # Lazy imports: keep module import cheap and let tests substitute these.
    from backend.pipeline.chunking import chunk_pages
    from backend.pipeline.embedding import index_chunks

    basename = (
        os.path.basename(original_filename)
        if original_filename
        else os.path.basename(file_path)
    )
    pages = _read_pages(file_path)
    if original_filename:
        # Pages were read from a staging path; cite the real document name.
        for page in pages:
            page["source_file"] = basename
    tables_found = sum(len(page.get("tables") or []) for page in pages)

    chunks = chunk_pages(pages)
    chunks_indexed = index_chunks(chunks, project_id, source_type=source_type)

    requirements_extracted = 0
    if source_type == "tender":
        requirements = extraction.extract_requirements(chunks, project_id, db_session)
        requirements_extracted = len(requirements)
    elif source_type == "change_request":
        requirements = extraction.process_change_request(chunks, project_id, db_session)
        requirements_extracted = len(requirements)
    elif source_type == "inspection_report":
        findings = _extract_inspection_findings(chunks, project_id, db_session)
        logger.info(
            "Inspection report %s produced %d finding(s)", basename, len(findings)
        )
    else:  # meeting_notes -> retrieval indexing only
        logger.info("Meeting notes %s indexed for retrieval only", basename)

    try:
        size_bytes = os.path.getsize(file_path)
    except OSError:
        size_bytes = 0

    document = IngestedDocument(
        project_id=project_id,
        filename=basename,
        source_type=source_type,
        pages=len(pages),
        chunks=chunks_indexed,
        tables_found=tables_found,
        size_bytes=size_bytes,
    )
    db_session.add(document)
    db_session.commit()

    result = {
        "project_id": project_id,
        "filename": basename,
        "source_type": source_type,
        "pages_processed": len(pages),
        "chunks_indexed": chunks_indexed,
        "tables_found": tables_found,
        "requirements_extracted": requirements_extracted,
    }
    logger.info("Ingested %s: %s", basename, result)
    return result
